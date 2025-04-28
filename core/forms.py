from django import forms
from .models import Document
import os

from django import forms
from .models import Document
import os
class DocumentForm(forms.ModelForm):
    class Meta:
        model = Document
        fields = ['title', 'file', 'printer']

    def clean(self):
        cleaned_data = super().clean()
        printer = cleaned_data.get('printer')
        file = cleaned_data.get('file')

        if not file or not printer:
            raise forms.ValidationError("Необходимо выбрать файл и принтер.")

        # Определение расширения файла
        extension = os.path.splitext(file.name)[1].lower()

        # Подсчет страниц
        pages = 0
        try:
            if extension == '.pdf':
                pages = self.get_pdf_page_count(file)
            # elif extension == '.docx':
            #     pages = self.get_docx_page_count(file)
            elif extension == '.pptx':
                pages = self.get_pptx_page_count(file)
            else:
                pages = 1
                pass
                # raise forms.ValidationError("Неподдерживаемый формат файла.")
        except Exception as e:
            # raise forms.ValidationError(f"Ошибка при обработке файла: {str(e)}")
            pass
        # Проверка количества бумаги
        if printer and printer.paper_count < pages:
            raise forms.ValidationError(
                f"Недостаточно бумаги для печати!!!. Нужно {pages} листов, а в наличии {printer.paper_count}."
            )

        cleaned_data['pages'] = pages
        return cleaned_data

    def get_pdf_page_count(self, file):
        # Пример подсчета страниц для PDF
        from PyPDF2 import PdfReader
        reader = PdfReader(file)
        return len(reader.pages)

    def get_docx_page_count(self, file):
        # Пример подсчета страниц для DOCX
        from docx import Document as DocxDocument
        document = DocxDocument(file)
        return len(document.paragraphs)  # Это пример, настройте под ваши нужды

    def get_pptx_page_count(self, file):
        # Пример подсчета страниц для PPTX
        from pptx import Presentation
        presentation = Presentation(file)
        return len(presentation.slides)

# class DocumentForm(forms.ModelForm):
#     class Meta:
#         model = Document
#         fields = ['title', 'file', 'printer']

#     def clean(self):
#         cleaned_data = super().clean()
#         printer = cleaned_data.get('printer')
#         file = cleaned_data.get('file')

#         if not file or not printer:
#             raise forms.ValidationError("Необходимо выбрать файл и принтер.")

#         # Определение расширения файла
#         extension = os.path.splitext(file.name)[1].lower()

#         # В зависимости от типа файла вызываем соответствующий метод подсчета страниц
#         document = Document(file=file, printer=printer)

#         try:
#             if extension == '.pdf':
#                 document.pages = document.get_pdf_page_count(file)
#             elif extension == '.docx':
#                 document.pages = document.get_docx_page_count(file)
#             elif extension == '.pptx':
#                 document.pages = document.get_pptx_page_count(file)
#             else:
#                 raise forms.ValidationError("Неподдерживаемый формат файла.")
#         except Exception as e:
#             raise forms.ValidationError(f"Ошибка при обработке файла: {str(e)}")

#         # Проверка количества бумаги
#         if printer.paper_count < document.pages:
#             raise forms.ValidationError(
#                 f"Недостаточно бумаги для печати. Нужно {document.pages} листов, а в наличии {printer.paper_count}."
#             )

#         return cleaned_data

